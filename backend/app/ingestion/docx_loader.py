"""DOCX text extraction using python-docx."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from app.ingestion.base import BaseDocumentLoader
from app.ingestion.exceptions import DocumentLoadError, EmptyDocumentError
from app.ingestion.models import ExtractedDocument, ExtractedPage

class DOCXDocumentLoader(BaseDocumentLoader):
    supported_extensions = {".docx"}

    def load(self, path: Path) -> ExtractedDocument:
        try:
            document = Document(path)
            blocks: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    blocks.append(text)

            for table_index, table in enumerate(document.tables, start=1):
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                if rows:
                    blocks.append(f"[Table {table_index}]\n" + "\n".join(rows))

            properties = document.core_properties
            metadata = {
                "title": properties.title or None,
                "subject": properties.subject or None,
                "author": properties.author or None,
                "keywords": properties.keywords or None,
                "comments": properties.comments or None,
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            }
        except Exception as exc:
            raise DocumentLoadError(f"Failed to read DOCX file '{path.name}'.") from exc

        text = "\n\n".join(blocks).strip()
        if not text:
            raise EmptyDocumentError(f"No extractable text was found in DOCX '{path.name}'.")

        return ExtractedDocument(
            source_path=str(path.resolve()),
            file_extension=path.suffix.lower(),
            pages=[ExtractedPage(page_number=1, text=text)],
            metadata=metadata,
        )
