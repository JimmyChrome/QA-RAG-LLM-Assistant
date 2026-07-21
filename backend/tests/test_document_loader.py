"""Tests for document text extraction."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from app.ingestion.exceptions import (
    DocumentFileNotFoundError,
    EmptyDocumentError,
    UnsupportedDocumentTypeError,
)
from app.ingestion.loader import DocumentLoader


@pytest.fixture
def loader() -> DocumentLoader:
    return DocumentLoader()


def test_load_txt(loader: DocumentLoader, tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("Quality assurance\nInternal assessment", encoding="utf-8")

    result = loader.load(path)

    assert result.file_extension == ".txt"
    assert result.page_count == 1
    assert "Quality assurance" in result.full_text
    assert result.metadata["encoding"] == "utf-8"


def test_load_markdown(loader: DocumentLoader, tmp_path: Path) -> None:
    path = tmp_path / "sample.md"
    path.write_text("# QA Policy\n\nThis is a policy.", encoding="utf-8")

    result = loader.load(path)

    assert result.file_extension == ".md"
    assert "# QA Policy" in result.full_text
    assert result.metadata["format"] == "markdown"


def test_load_docx_with_table(loader: DocumentLoader, tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("QA Manual", level=1)
    document.add_paragraph("This document explains QA procedures.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Area"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Assessment"
    table.cell(1, 1).text = "Complete"
    document.save(path)

    result = loader.load(path)

    assert result.file_extension == ".docx"
    assert "QA Manual" in result.full_text
    assert "Assessment | Complete" in result.full_text
    assert result.metadata["table_count"] == 1


def test_load_pdf_page_by_page(loader: DocumentLoader, tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    pdf = fitz.open()

    first_page = pdf.new_page()
    first_page.insert_text((72, 72), "Page one QA content")

    second_page = pdf.new_page()
    second_page.insert_text((72, 72), "Page two accreditation content")

    pdf.save(path)
    pdf.close()

    result = loader.load(path)

    assert result.file_extension == ".pdf"
    assert result.page_count == 2
    assert result.pages[0].page_number == 1
    assert "Page one QA content" in result.pages[0].text
    assert "Page two accreditation content" in result.pages[1].text


def test_reject_empty_document(loader: DocumentLoader, tmp_path: Path) -> None:
    path = tmp_path / "empty.txt"
    path.write_text("   \n\n", encoding="utf-8")

    with pytest.raises(EmptyDocumentError):
        loader.load(path)


def test_reject_unsupported_document(loader: DocumentLoader, tmp_path: Path) -> None:
    path = tmp_path / "sample.csv"
    path.write_text("name,value", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentTypeError):
        loader.load(path)


def test_reject_missing_document(loader: DocumentLoader, tmp_path: Path) -> None:
    with pytest.raises(DocumentFileNotFoundError):
        loader.load(tmp_path / "missing.pdf")
